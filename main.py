from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
import os, uuid
from io import BytesIO
from docx import Document
import mammoth
from bs4 import BeautifulSoup, NavigableString, Tag

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

templates = Jinja2Templates(directory="templates")
STORAGE = "storage"
os.makedirs(STORAGE, exist_ok=True)

documents = {}

@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request, "documents": documents})

@app.post("/documents")
def create_doc(name: str = Form(...), content: str = Form("")):
    doc_id = str(uuid.uuid4())
    documents[doc_id] = {"name": name, "raw_html": content}
    return RedirectResponse(url="/", status_code=303)

@app.get("/documents/{doc_id}", response_class=HTMLResponse)
def edit_doc(request: Request, doc_id: str):
    meta = documents.get(doc_id)
    if not meta:
        return HTMLResponse("Not found", status_code=404)
    return templates.TemplateResponse("index.html", {
        "request": request,
        "documents": documents,
        "edit_id": doc_id,
        "edit_name": meta["name"],
        "edit_content": meta.get("raw_html", "")
    })

@app.post("/documents/{doc_id}/update")
def update_doc(doc_id: str, name: str = Form(...), content: str = Form(...)):
    meta = documents.get(doc_id)
    if not meta:
        return HTMLResponse("Not found", status_code=404)
    meta["name"] = name
    meta["raw_html"] = content
    return RedirectResponse(url="/", status_code=303)

# Enhanced inline formatting: bold, italic, underline, strikethrough, nested tags
# You can extend for more as needed

def add_inline_content(p, element):
    for node in element.children:
        if isinstance(node, NavigableString):
            p.add_run(str(node))
        elif isinstance(node, Tag):
            tag = node.name.lower()
            if tag in ["strong", "b"]:
                r = p.add_run(node.get_text())
                r.bold = True
            elif tag in ["em", "i"]:
                r = p.add_run(node.get_text())
                r.italic = True
            elif tag == "u":
                r = p.add_run(node.get_text())
                r.underline = True
            elif tag in ["del", "s", "strike"]:
                r = p.add_run(node.get_text())
                r.font.strike = True
            elif tag == "span":
                # Handle inline styles like color or font-weight if present
                style = node.get("style", "").lower()
                run = p.add_run(node.get_text())
                if "bold" in style:
                    run.bold = True
                if "italic" in style:
                    run.italic = True
                if "underline" in style:
                    run.underline = True
                # color example
                if "color:" in style:
                    try:
                        from docx.shared import RGBColor
                        import re
                        m = re.search(r"color:\s*#([0-9a-fA-F]{6})", style)
                        if m:
                            hex_color = m.group(1)
                            run.font.color.rgb = RGBColor(int(hex_color[0:2],16),int(hex_color[2:4],16),int(hex_color[4:6],16))
                    except:
                        pass
            else:
                # recurse for nested tags
                add_inline_content(p, node)

@app.get("/documents/{doc_id}/download")
def download_doc(doc_id: str):
    meta = documents.get(doc_id)
    if not meta:
        return JSONResponse(status_code=404, content={"error": "not found"})
    html_content = meta.get("raw_html", "").strip()
    if not html_content:
        html_content = "<p>(No content)</p>"

    soup = BeautifulSoup(html_content, "html.parser")
    doc = Document()

    def process_element(el):
        tag = el.name.lower() if el.name else ""
        if tag.startswith("h") and len(tag) == 2 and tag[1].isdigit():
            p = doc.add_paragraph()
            run = p.add_run(el.get_text())
            run.bold = True
            run.font.size = None  # Could map to heading sizes if needed
        elif tag == "p":
            p = doc.add_paragraph()
            add_inline_content(p, el)
        elif tag == "ul":
            for li in el.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                add_inline_content(p, li)
        elif tag == "ol":
            for li in el.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number')
                add_inline_content(p, li)
        elif tag == "table":
            rows = el.find_all('tr', recursive=False)
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td','th'], recursive=False)))
                for r_idx,row in enumerate(rows):
                    cells = row.find_all(['td','th'], recursive=False)
                    for c_idx, cell in enumerate(cells):
                        t.cell(r_idx,c_idx).text = cell.get_text(strip=True)

    for elem in soup.find_all(recursive=False):
        process_element(elem)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = meta["name"]
    if not filename.endswith(".docx"):
        filename += ".docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/documents/{doc_id}/delete")
def delete_doc(doc_id: str):
    documents.pop(doc_id, None)
    return RedirectResponse(url="/", status_code=303)

@app.post("/upload")
def upload_doc(file: UploadFile = File(...)):
    doc_id = str(uuid.uuid4())
    path = os.path.join(STORAGE, f"{doc_id}.docx")
    content = file.file.read()
    with open(path, "wb") as f:
        f.write(content)
    try:
        with open(path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
    except Exception:
        html = "<p>(Unable to parse uploaded document)</p>"
    documents[doc_id] = {"name": file.filename, "raw_html": html}
    return RedirectResponse(url="/", status_code=303)
